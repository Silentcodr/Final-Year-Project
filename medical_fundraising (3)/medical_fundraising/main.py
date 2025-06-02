from flask import Flask
from flask import Flask, render_template, Response, redirect, request, session, abort, url_for, jsonify
import os
import time
import datetime
from random import randint
import cv2
import PIL.Image
from PIL import Image
import imagehash
import hashlib
import PyPDF2

#pip install PyMuPDF
import fitz
import docx
from docx import Document
from diff_match_patch import diff_match_patch
#from pdf2docx import parse
from typing import Tuple
from PIL import Image, ImageDraw, ImageFont
import textwrap
from spire.doc import *
from spire.doc.common import *
import pytesseract
from skimage.metrics import structural_similarity
#from fuzzywuzzy import fuzz

from flask import send_file
from werkzeug.utils import secure_filename
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import mysql.connector

mydb = mysql.connector.connect(
  host="localhost",
  user="root",
  passwd="",
  charset="utf8",
  database="medical_fundraising"
)
app = Flask(__name__)
##session key
app.secret_key = 'abcdef'
#######
UPLOAD_FOLDER = 'upload'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
#####

@app.route('/',methods=['POST','GET'])
def index():
    act=""
    msg=""

    #now1 = datetime.datetime.now()
    #rtime=now1.strftime("%H:%M")
    #print(rtime)

    return render_template('index.html',msg=msg,act=act)

@app.route('/login_admin', methods=['GET', 'POST'])
def login_admin():
    msg=""
    
    if request.method=='POST':
        uname=request.form['uname']
        pwd=request.form['pass']
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM mf_admin WHERE username = %s AND password = %s', (uname, pwd))
        account = cursor.fetchone()
        if account:
            session['username'] = uname
            
            return redirect(url_for('admin'))
        else:
            msg = 'Incorrect username/password!'
    return render_template('login_admin.html',msg=msg)

@app.route('/login_donator', methods=['GET', 'POST'])
def login_donator():
    msg=""
    
    if request.method=='POST':
        uname=request.form['uname']
        pwd=request.form['pass']
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM mf_donator WHERE uname = %s AND pass = %s', (uname, pwd))
        account = cursor.fetchone()
        if account:
            session['username'] = uname
           
            return redirect(url_for('dr_home'))
        else:
            msg = 'Incorrect username/password!'
    return render_template('login_donator.html',msg=msg)

@app.route('/login', methods=['GET', 'POST'])
def login():
    msg=""
    
    if request.method=='POST':
        uname=request.form['uname']
        pwd=request.form['pass']
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM mf_user WHERE uname = %s AND pass = %s', (uname, pwd))
        account = cursor.fetchone()
        if account:
            session['username'] = uname
            return redirect(url_for('userhome'))
        else:
            msg = 'Incorrect username/password!'
    return render_template('login.html',msg=msg)



@app.route('/reg_user',methods=['POST','GET'])
def reg_user():
    msg=""
    act=""
    if request.method=='POST':
        name=request.form['name']
        mobile=request.form['mobile']
        email=request.form['email']
        address=request.form['address']
        city=request.form['city']
        acc_name=request.form['acc_name']
        bank_name=request.form['bank_name']
        account=request.form['account']
        branch=request.form['branch']
        gpay_number=request.form['gpay_number']
        uname=request.form['uname']
        pass1=request.form['pass']
      
        now = datetime.datetime.now()
        rdate=now.strftime("%d-%m-%Y")
        mycursor = mydb.cursor()

        mycursor.execute("SELECT count(*) FROM mf_user where uname=%s",(uname, ))
        cnt = mycursor.fetchone()[0]
        if cnt==0:
            mycursor.execute("SELECT max(id)+1 FROM mf_user")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1
            sql = "INSERT INTO mf_user(id,name,mobile,email,address,city,acc_name,bank_name,account,branch,gpay_number,uname,pass) VALUES (%s, %s, %s, %s, %s, %s, %s,%s,%s,%s,%s,%s,%s)"
            val = (maxid,name,mobile,email,address,city,acc_name,bank_name,account,branch,gpay_number,uname,pass1)
            print(sql)
            mycursor.execute(sql, val)
            mydb.commit()            
            print(mycursor.rowcount, "record inserted.")
            msg='success'
            
            #if mycursor.rowcount==1:
            #    result="Registered Success"
            
        else:
            msg="fail"
    return render_template('reg_user.html',msg=msg)

@app.route('/reg_donator',methods=['POST','GET'])
def reg_donator():
    msg=""
    act=""
    mycursor = mydb.cursor()
    
    
    if request.method=='POST':
        
        name=request.form['name']
        mobile=request.form['mobile']
        email=request.form['email']
        
        city=request.form['city']
        description=request.form['description']
        uname=request.form['uname']
        pass1=request.form['pass']
      
        now = datetime.datetime.now()
        rdate=now.strftime("%d-%m-%Y")
        mycursor = mydb.cursor()

        mycursor.execute("SELECT count(*) FROM mf_donator where uname=%s",(uname, ))
        cnt = mycursor.fetchone()[0]
        if cnt==0:
            mycursor.execute("SELECT max(id)+1 FROM mf_donator")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1
            sql = "INSERT INTO mf_donator(id,name,mobile,email,city,description,uname,pass) VALUES (%s, %s, %s, %s, %s, %s, %s,%s)"
            val = (maxid,name,mobile,email,city,description,uname,pass1)
            print(sql)
            mycursor.execute(sql, val)
            mydb.commit()            
            print(mycursor.rowcount, "record inserted.")
            msg='success'
            
            #if mycursor.rowcount==1:
            #    result="Registered Success"
            
        else:
            msg="fail"
     
    return render_template('reg_donator.html',msg=msg)



@app.route('/userhome',methods=['POST','GET'])
def userhome():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where uname=%s && req_status=1 order by id desc",(uname,))
    data2 = mycursor.fetchall()
    
    return render_template('web/userhome.html',msg=msg,data=data,data2=data2)

def calculate_hash(file_path):
    # Calculate the hash value of a file
    hasher = hashlib.sha256()
    with open(file_path, 'rb') as f:
        while True:
            data = f.read(65536)  # Read the file in chunks to avoid loading it entirely into memory
            if not data:
                break
            hasher.update(data)
    return hasher.hexdigest()

def hash_difference_percentage(hash1, hash2):
    # Calculate the percentage difference between two hash values
    if len(hash1) != len(hash2):
        raise ValueError("Hash values must have the same length")
    
    difference_count = sum(c1 != c2 for c1, c2 in zip(hash1, hash2))
    total_length = len(hash1)
    percentage_difference = (difference_count / total_length) * 100
    return percentage_difference


def sha256_hash(file_path):
    try:
        with open(file_path, 'rb') as file:
            # Read the entire file
            data = file.read()
            # Calculate the SHA-256 hash
            sha256_hash = hashlib.sha256(data).hexdigest()
            return sha256_hash
    except FileNotFoundError:
        print("File not found")
        return None
    
@app.route('/add_post',methods=['POST','GET'])
def add_post():
    msg=""
    uname=""
    pid=request.args.get("pid")
    act=request.args.get("act")
    if 'username' in session:
        uname = session['username']
    sid=0 
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_treatment")
    data1 = mycursor.fetchall()

    mycursor.execute("SELECT * FROM mf_status")
    data2 = mycursor.fetchall()
    

    now = datetime.datetime.now()
    rdate=now.strftime("%d-%m-%Y")

    if request.method=='POST':
        pat_name=request.form['pat_name']
        gender=request.form['gender']
        dob=request.form['dob']
        address=request.form['address']
        city=request.form['city']
        aadhar=request.form['aadhar']
        hospital=request.form['hospital']
        location=request.form['location']
        hos_city=request.form['hos_city']
        patient_id=request.form['patient_id']
        treatment=request.form['treatment']
        hospital_status=request.form['hospital_status']
        
        req_amount=request.form['req_amount']

        treatment1=""
        if treatment=="other":
            treatment1=request.form['oth_treatment']
            mycursor.execute("SELECT max(id)+1 FROM mf_treatment")
            maxid2 = mycursor.fetchone()[0]
            if maxid2 is None:
                maxid2=1
            sql2 = "INSERT INTO mf_treatment(id,treatment) VALUES (%s,%s)"
            val2 = (maxid2,treatment1)
            mycursor.execute(sql2, val2)
            mydb.commit()
            
        else:
            treatment1=treatment

        hos_status=""
        if hospital_status=="other":
            hos_status=request.form['oth_status']
            mycursor.execute("SELECT max(id)+1 FROM mf_status")
            maxid3 = mycursor.fetchone()[0]
            if maxid3 is None:
                maxid3=1
            sql3 = "INSERT INTO mf_status(id,status) VALUES (%s,%s)"
            val3 = (maxid3,hos_status)
            mycursor.execute(sql3, val3)
            mydb.commit()
        else:
            hos_status=hospital_status
        
        mycursor.execute("SELECT max(id)+1 FROM mf_post")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1

        ds=dob.split("-")
        dob1=ds[2]+"-"+ds[1]+"-"+ds[0]

        mycursor.execute("SELECT * FROM mf_patient_data")
        d2 = mycursor.fetchall()

        idd=0
        for dd in d2:
            if aadhar==dd[6] and hospital==dd[7]:
                if pat_name==dd[1] or patient_id==dd[10]:
                    idd=dd[0]
                    break
            else:
                idd=0
        print("idd")
        print(idd)
        if idd>0:

            mycursor.execute("SELECT * FROM mf_patient_data where id=%s",(idd,))
            d3 = mycursor.fetchone()
            bamount=d3[13]

            bill=int(bamount)
            ramount=int(req_amount)
            if ramount<=bill:
            
                sql = "INSERT INTO mf_post(id,uname,pat_name,gender,dob,address,city,aadhar,hospital,location,hos_city,patient_id,treatment,hospital_status,req_amount,req_date,req_status,sid) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
                val = (maxid,uname,pat_name,gender,dob1,address,city,aadhar,hospital,location,hos_city,patient_id,treatment1,hos_status,req_amount,rdate,'0',idd)
                mycursor.execute(sql, val)
                mydb.commit()

                pid=str(maxid)

                mycursor.execute("SELECT count(*) FROM mf_files where post_id=%s",(pid,))
                cnt = mycursor.fetchone()[0]
                if cnt>0:
                    mycursor.execute("update mf_files set post_id=%s,status=1 where status=0",(pid,))
                    mydb.commit()
                    msg="ok"
                else:
                    msg="nofile"
            else:
                msg="amt"

        else:
            msg="fail"

    x=0
    if act=="yes":

        mycursor.execute("SELECT count(*) FROM mf_files where post_id=%s",(pid,))
        cnt = mycursor.fetchone()[0]
        if cnt>0:
            mycursor.execute("SELECT * FROM mf_files where post_id=%s",(pid,))
            fd1 = mycursor.fetchall()
            for fd11 in fd1:
                sid=fd11[7]
                if sid>0:
                    if str(sid)==pid:
                        x+=1
                else:
                    x+=1
            if x>0:    
                
                msg="ok"
            else:
                msg="fail"
        else:
            msg="nofile"
   
    #text1 = extract_text_from_docx("sample/docmm.docx")
    #text2 = extract_text_from_docx("sample/docmm2.docx")
    #differences = highlight_differences(text1, text2)
                
    return render_template('web/add_post.html',msg=msg,act=act,data=data,data1=data1,data2=data2,pid=pid)

@app.route('/verify1',methods=['POST','GET'])
def verify1():
    msg=""
    uname=""
    st=""
    act=request.args.get("act")
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    return render_template('web/verify1.html',msg=msg,act=act,data=data,pid=pid)

def convert_pdf2docx(input_file: str, output_file: str, pages: Tuple = None):
    """Converts pdf to docx"""
    if pages:
        pages = [int(i) for i in list(pages) if i.isnumeric()]
    result = parse(pdf_file=input_file,
                   docx_with_path=output_file, pages=pages)
    summary = {
        "File": input_file, "Pages": str(pages), "Output File": output_file
    }
    # Printing Summary
    print("## Summary ########################################################")
    print("\n".join("{}:{}".format(i, j) for i, j in summary.items()))
    print("###################################################################")
    return result

def word_to_img(wfile,fid):
    # Create a Document object
    document = Document()
    # Load a Word DOCX file
    document.LoadFromFile("static/upload/"+wfile)
    # Or load a Word DOC file
    #document.LoadFromFile("Sample.doc")

    # Convert the document to a list of image streams
    image_streams = document.SaveImageToStreams(ImageType.Bitmap)

    # Incremental counter
    i = 1

    # Save each image stream to a PNG file
    for image in image_streams:
        image_name = "m"+str(fid)+"_"+str(i) + ".png"
        with open("static/upload/"+image_name,'wb') as image_file:
            image_file.write(image.ToArray())
        i += 1

    # Close the document
    document.Close()
    return i

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    return full_text

def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    return full_text

def highlight_differences(text1, text2):
    dmp = diff_match_patch()
    diffs = dmp.diff_main(text1, text2)
    dmp.diff_cleanupSemantic(diffs)
    diff_html = dmp.diff_prettyHtml(diffs)
    return diff_html

@app.route('/add_attach', methods=['GET', 'POST'])
def add_attach():
    msg=""
    st=""
    data2=[]
    view=""
    fname=""
    textpdf=""
    textdoc=""
    ccode=""
    tt=100
    pid=request.args.get("pid")
    nn=0
    sid=0
    per=0
    x=0
    uname=""
    if 'username' in session:
        uname = session['username']

    ss='0'
    if pid=='0':
        ss='0'
    else:
        ss='1'
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()
    
    if request.method=='POST':
        detail=request.form['detail']
        file = request.files['file']

        mycursor.execute("SELECT max(id)+1 FROM mf_files")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1

        fn1=file.filename
        fn="f"+str(maxid)+fn1                
        file.save(os.path.join("static/upload", fn))

        mycursor.execute("SELECT * FROM mf_patient_files")
        fd1 = mycursor.fetchall()
        for fd11 in fd1:
            hash1 = calculate_hash("static/upload/"+fn)
            hash2 = calculate_hash("static/data/"+fd11[2])
            percentage_difference = hash_difference_percentage(hash1, hash2)
            per=tt-percentage_difference
            print("percent")
            print(per)
            if per==100:
                x+=1
                sid=fd11[3]
                break

        
            
        sql = "INSERT INTO mf_files(id,uname,detail,filename,post_id,status,sid) VALUES (%s, %s, %s, %s,%s,%s,%s)"
        val = (maxid,uname,detail,fn,pid,ss,sid)
        mycursor.execute(sql, val)
        mydb.commit()
        
        fname=fn
        ex1=fname.split(".")
        if ex1[1]=="pdf":
            p1=fname.split(".")
            fname2=p1[0]+".docx"
            convert_pdf2docx("static/upload/"+fname,"static/upload/"+fname2)
            nn=word_to_img(fname2,maxid)
            nn1=nn-1
            mycursor.execute("update mf_files set img_count=%s where id=%s",(nn,maxid))
            mydb.commit()
        elif ex1[1]=="docx":
            nn=word_to_img(fname,maxid)
            nn1=nn-1
            mycursor.execute("update mf_files set img_count=%s where id=%s",(nn1,maxid))
            mydb.commit()
          
        msg="success"

    mycursor.execute("SELECT * FROM mf_files where uname=%s && post_id=%s",(uname,pid))
    data1 = mycursor.fetchall()

    mycursor.execute("SELECT count(*) FROM mf_files where uname=%s && post_id=%s",(uname,pid))
    cnt = mycursor.fetchone()[0]
    if cnt>0:
        st="1"
        
        mycursor.execute("SELECT * FROM mf_files where uname=%s && post_id=%s",(uname,pid))
        d1 = mycursor.fetchall()
        for d2 in d1:
            dt=[]
            dt.append(d2[0])
            dt.append(d2[1])
            dt.append(d2[2])
            dt.append(d2[3])
            dt.append(d2[4])
            dt.append(d2[5])
            dt.append(d2[6])

            #7#
            ex=d2[3].split(".")
            if ex[1]=="png":
                dt.append("png")
            elif ex[1]=="jpg":
                dt.append("jpg")
            elif ex[1]=="jpeg":
                dt.append("jpeg")
            elif ex[1]=="pdf":
                dt.append("pdf")
                fname=d2[3]
                file2="static/upload/"+fname
                textpdf = extract_text_from_pdf(file2)
               
            
            elif ex[1]=="docx":
                dt.append("docx")
                fname=d2[3]
                file2="static/upload/"+fname
    
                textdoc = extract_text_from_docx(file2)
                
            
            elif ex[1]=="txt":
                dt.append("txt")
                fname=d2[3]
                file1 = open("static/upload/"+fname, 'r')
                Lines = file1.readlines()
                 
                count = 0
                result=""
                # Strips the newline character
                for line in Lines:
                    result = "".join(line for line in Lines if not line.isspace())
                    count += 1
                    #print("Line{}: {}".format(count, line.strip()))
                ccode=result
            
            else:
                dt.append("")

            #8#
            if d2[6]>0:
                dt.append("yes")
            else:
                dt.append("no")
            #9#
            if d2[6]>0:
                i=1
                dtt=[]
                
                while i<=d2[6]:
                    mg="m"+str(d2[0])+"_"+str(i)+".png"
                    dtt.append(mg)
                    i+=1
                dt.append(dtt)
            ##
            data2.append(dt)
           

    return render_template('web/add_attach.html',msg=msg,data=data,data2=data2,st=st,ccode=ccode,textpdf=textpdf,textdoc=textdoc)

#YoloV8 - Object Detection
def object_detect():
    ap = argparse.ArgumentParser()
    ap.add_argument('-i', '--image', required=True,
                    help = 'path to input image')
    ap.add_argument('-c', '--config', required=True,
                    help = 'path to yolo config file')
    ap.add_argument('-w', '--weights', required=True,
                    help = 'path to yolo pre-trained weights')
    ap.add_argument('-cl', '--classes', required=True,
                    help = 'path to text file containing class names')
    args = ap.parse_args()

def get_output_layers(net):
    
    layer_names = net.getLayerNames()
    try:
        output_layers = [layer_names[i - 1] for i in net.getUnconnectedOutLayers()]
    except:
        output_layers = [layer_names[i[0] - 1] for i in net.getUnconnectedOutLayers()]
    return output_layers

def draw_prediction(img, class_id, confidence, x, y, x_plus_w, y_plus_h):
    label = str(classes[class_id])
    color = COLORS[class_id]
    cv2.rectangle(img, (x,y), (x_plus_w,y_plus_h), color, 2)
    cv2.putText(img, label, (x-10,y-10), cv2.FONT_HERSHEY_SIMPLEX, 0.5, color, 2)

    
    image = cv2.imread(args.image)
    Width = image.shape[1]
    Height = image.shape[0]
    scale = 0.00392
    classes = None
    with open(args.classes, 'r') as f:
        classes = [line.strip() for line in f.readlines()]

    COLORS = np.random.uniform(0, 255, size=(len(classes), 3))
    net = cv2.dnn.readNet(args.weights, args.config)
    blob = cv2.dnn.blobFromImage(image, scale, (416,416), (0,0,0), True, crop=False)
    net.setInput(blob)
    outs = net.forward(get_output_layers(net))

    class_ids = []
    confidences = []
    boxes = []
    conf_threshold = 0.5
    nms_threshold = 0.4

    for out in outs:
        for detection in out:
            scores = detection[5:]
            class_id = np.argmax(scores)
            confidence = scores[class_id]
            if confidence > 0.5:
                center_x = int(detection[0] * Width)
                center_y = int(detection[1] * Height)
                w = int(detection[2] * Width)
                h = int(detection[3] * Height)
                x = center_x - w / 2
                y = center_y - h / 2
                class_ids.append(class_id)
                confidences.append(float(confidence))
                boxes.append([x, y, w, h])


    indices = cv2.dnn.NMSBoxes(boxes, confidences, conf_threshold, nms_threshold)
    for i in indices:
        try:
            box = boxes[i]
        except:
            i = i[0]
            box = boxes[i]
        
        x = box[0]
        y = box[1]
        w = box[2]
        h = box[3]
        draw_prediction(image, class_ids[i], confidences[i], round(x), round(y), round(x+w), round(y+h))
    
#PaddleOCR
def PaddleOCR():
   
        params = parse_args(mMain=False)
        params.__dict__.update(**kwargs)
        assert (
            params.ocr_version in SUPPORT_OCR_MODEL_VERSION
        ), "ocr_version must in {}, but get {}".format(
            SUPPORT_OCR_MODEL_VERSION, params.ocr_version
        )
        params.use_gpu = check_gpu(params.use_gpu)

        if not params.show_log:
            logger.setLevel(logging.INFO)
        self.use_angle_cls = params.use_angle_cls
        lang, det_lang = parse_lang(params.lang)

        # init model dir
        det_model_config = get_model_config("OCR", params.ocr_version, "det", det_lang)
        params.det_model_dir, det_url = confirm_model_dir_url(
            params.det_model_dir,
            os.path.join(BASE_DIR, "whl", "det", det_lang),
            det_model_config["url"],
        )
        rec_model_config = get_model_config("OCR", params.ocr_version, "rec", lang)
        params.rec_model_dir, rec_url = confirm_model_dir_url(
            params.rec_model_dir,
            os.path.join(BASE_DIR, "whl", "rec", lang),
            rec_model_config["url"],
        )
        cls_model_config = get_model_config("OCR", params.ocr_version, "cls", "ch")
        params.cls_model_dir, cls_url = confirm_model_dir_url(
            params.cls_model_dir,
            os.path.join(BASE_DIR, "whl", "cls"),
            cls_model_config["url"],
        )
        if params.ocr_version in ["PP-OCRv3", "PP-OCRv4"]:
            params.rec_image_shape = "3, 48, 320"
        else:
            params.rec_image_shape = "3, 32, 320"
        if kwargs.get("rec_image_shape") is not None:
            params.rec_image_shape = kwargs.get("rec_image_shape")
        # download model if using paddle infer
        if not params.use_onnx:
            maybe_download(params.det_model_dir, det_url)
            maybe_download(params.rec_model_dir, rec_url)
            maybe_download(params.cls_model_dir, cls_url)

        if params.det_algorithm not in SUPPORT_DET_MODEL:
            logger.error("det_algorithm must in {}".format(SUPPORT_DET_MODEL))
            sys.exit(0)
        if params.rec_algorithm not in SUPPORT_REC_MODEL:
            logger.error("rec_algorithm must in {}".format(SUPPORT_REC_MODEL))
            sys.exit(0)

        if params.rec_char_dict_path is None:
            params.rec_char_dict_path = str(
                Path(__file__).parent / rec_model_config["dict_path"]
            )

        logger.debug(params)
        # init det_model and rec_model
        super().__init__(params)
        self.page_num = params.page_num

        assert isinstance(img, (np.ndarray, list, str, bytes))
        if isinstance(img, list) and det == True:
            logger.error("When input a list of images, det must be false")
            exit(0)
        if cls == True and self.use_angle_cls == False:
            logger.warning(
                "Since the angle classifier is not initialized, it will not be used during the forward process"
            )

        img, flag_gif, flag_pdf = check_img(img, alpha_color)
        # for infer pdf file
        if isinstance(img, list) and flag_pdf:
            if self.page_num > len(img) or self.page_num == 0:
                imgs = img
            else:
                imgs = img[: self.page_num]
        else:
            imgs = [img]

        def preprocess_image(_image):
            _image = alpha_to_color(_image, alpha_color)
            if inv:
                _image = cv2.bitwise_not(_image)
            if bin:
                _image = binarize_img(_image)
            return _image

        if det and rec:
            ocr_res = []
            for img in imgs:
                img = preprocess_image(img)
                dt_boxes, rec_res, _ = self.__call__(img, cls, slice)
                if not dt_boxes and not rec_res:
                    ocr_res.append(None)
                    continue
                tmp_res = [[box.tolist(), res] for box, res in zip(dt_boxes, rec_res)]
                ocr_res.append(tmp_res)
            return ocr_res
        elif det and not rec:
            ocr_res = []
            for img in imgs:
                img = preprocess_image(img)
                dt_boxes, elapse = self.text_detector(img)
                if dt_boxes.size == 0:
                    ocr_res.append(None)
                    continue
                tmp_res = [box.tolist() for box in dt_boxes]
                ocr_res.append(tmp_res)
            return ocr_res
        else:
            ocr_res = []
            cls_res = []
            for img in imgs:
                if not isinstance(img, list):
                    img = preprocess_image(img)
                    img = [img]
                if self.use_angle_cls and cls:
                    img, cls_res_tmp, elapse = self.text_classifier(img)
                    if not rec:
                        cls_res.append(cls_res_tmp)
                rec_res, elapse = self.text_recognizer(img)
                ocr_res.append(rec_res)
            if not rec:
                return cls_res
            return ocr_res


def PPStructure():
   
        params = parse_args(mMain=False)
        params.__dict__.update(**kwargs)
        assert (
            params.structure_version in SUPPORT_STRUCTURE_MODEL_VERSION
        ), "structure_version must in {}, but get {}".format(
            SUPPORT_STRUCTURE_MODEL_VERSION, params.structure_version
        )
        params.use_gpu = check_gpu(params.use_gpu)
        params.mode = "structure"

        if not params.show_log:
            logger.setLevel(logging.INFO)
        lang, det_lang = parse_lang(params.lang)
        if lang == "ch":
            table_lang = "ch"
        else:
            table_lang = "en"
        if params.structure_version == "PP-Structure":
            params.merge_no_span_structure = False

        # init model dir
        det_model_config = get_model_config("OCR", params.ocr_version, "det", det_lang)
        params.det_model_dir, det_url = confirm_model_dir_url(
            params.det_model_dir,
            os.path.join(BASE_DIR, "whl", "det", det_lang),
            det_model_config["url"],
        )
        rec_model_config = get_model_config("OCR", params.ocr_version, "rec", lang)
        params.rec_model_dir, rec_url = confirm_model_dir_url(
            params.rec_model_dir,
            os.path.join(BASE_DIR, "whl", "rec", lang),
            rec_model_config["url"],
        )
        table_model_config = get_model_config(
            "STRUCTURE", params.structure_version, "table", table_lang
        )
        params.table_model_dir, table_url = confirm_model_dir_url(
            params.table_model_dir,
            os.path.join(BASE_DIR, "whl", "table"),
            table_model_config["url"],
        )
        layout_model_config = get_model_config(
            "STRUCTURE", params.structure_version, "layout", lang
        )
        params.layout_model_dir, layout_url = confirm_model_dir_url(
            params.layout_model_dir,
            os.path.join(BASE_DIR, "whl", "layout"),
            layout_model_config["url"],
        )
        formula_model_config = get_model_config(
            "STRUCTURE", params.structure_version, "formula", lang
        )
        params.formula_model_dir, formula_url = confirm_model_dir_url(
            params.formula_model_dir,
            os.path.join(BASE_DIR, "whl", "formula"),
            formula_model_config["url"],
        )
        # download model
        if not params.use_onnx:
            maybe_download(params.det_model_dir, det_url)
            maybe_download(params.rec_model_dir, rec_url)
            maybe_download(params.table_model_dir, table_url)
            maybe_download(params.layout_model_dir, layout_url)
            maybe_download(params.formula_model_dir, formula_url)

        if params.rec_char_dict_path is None:
            params.rec_char_dict_path = str(
                Path(__file__).parent / rec_model_config["dict_path"]
            )
        if params.table_char_dict_path is None:
            params.table_char_dict_path = str(
                Path(__file__).parent / table_model_config["dict_path"]
            )
        if params.layout_dict_path is None:
            params.layout_dict_path = str(
                Path(__file__).parent / layout_model_config["dict_path"]
            )
        if params.formula_char_dict_path is None:
            params.formula_char_dict_path = str(
                Path(__file__).parent / formula_model_config["dict_path"]
            )
        logger.debug(params)
        super().__init__(params)
#Fuzzywuzzy
def test_Fuzzywuzzy():
    strings = ["atlanta braves", "Cães danados",
               "//// Mets $$$", "Ça va?"]
    for string in strings:
        proc_string = StringProcessor.replace_non_letters_non_numbers_with_whitespace(string)
        regex = re.compile(r"(?ui)[\W]")
        for expr in regex.finditer(proc_string):
            self.assertEqual(expr.group(), " ")

    def test_dont_condense_whitespace(self):
        s1 = "new york mets - atlanta braves"
        s2 = "new york mets atlanta braves"
        p1 = StringProcessor.replace_non_letters_non_numbers_with_whitespace(s1)
        p2 = StringProcessor.replace_non_letters_non_numbers_with_whitespace(s2)
        self.assertNotEqual(p1, p2)


    def setUp(self):
        self.s1 = "new"
        self.s1a = "mets"
        self.s2 = "new YORK mets"
        self.s3 = "the wonderful new york mets"
        self.s4 = "new york mets vs atlanta braves"
        self.s5 = "atlanta braves vs new york mets"
        self.s6 = "new york mets - atlanta braves"
        self.mixed_strings = [
            "Lorem Ipsum is simply dummy text of the printing and typesetting industry.",
            "C'est la vie",
            "Ça va?",
            "Cães danados",
            "\xacCamarões assados",
            "a\xac\u1234\u20ac\U00008000",
            "\u00C1"
        ]

    def tearDown(self):
        pass

    def test_asciidammit(self):
        for s in self.mixed_strings:
            utils.asciidammit(s)

    def test_asciionly(self):
        for s in self.mixed_strings:
            # ascii only only runs on strings
            s = utils.asciidammit(s)
            utils.asciionly(s)

    def test_fullProcess(self):
        for s in self.mixed_strings:
            utils.full_process(s)

    def test_fullProcessForceAscii(self):
        for s in self.mixed_strings:
            utils.full_process(s, force_ascii=True)

        self.cirque_strings = [
            "cirque du soleil - zarkana - las vegas",
            "cirque du soleil ",
            "cirque du soleil las vegas",
            "zarkana las vegas",
            "las vegas cirque du soleil at the bellagio",
            "zarakana - cirque du soleil - bellagio"
        ]

        self.baseball_strings = [
            "new york mets vs chicago cubs",
            "chicago cubs vs chicago white sox",
            "philladelphia phillies vs atlanta braves",
            "braves vs mets",
        ]

    def tearDown(self):
        pass

    def testEqual(self):
        self.assertEqual(fuzz.ratio(self.s1, self.s1a), 100)
        self.assertEqual(fuzz.ratio(self.s8, self.s8a), 100)
        self.assertEqual(fuzz.ratio(self.s9, self.s9a), 100)

    def testCaseInsensitive(self):
        self.assertNotEqual(fuzz.ratio(self.s1, self.s2), 100)
        self.assertEqual(fuzz.ratio(utils.full_process(self.s1), utils.full_process(self.s2)), 100)

    def testPartialRatio(self):
        self.assertEqual(fuzz.partial_ratio(self.s1, self.s3), 100)

    def testTokenSortRatio(self):
        self.assertEqual(fuzz.token_sort_ratio(self.s1, self.s1a), 100)

    def testPartialTokenSortRatio(self):
        self.assertEqual(fuzz.partial_token_sort_ratio(self.s1, self.s1a), 100)
        self.assertEqual(fuzz.partial_token_sort_ratio(self.s4, self.s5), 100)
        self.assertEqual(fuzz.partial_token_sort_ratio(self.s8, self.s8a, full_process=False), 100)
        self.assertEqual(fuzz.partial_token_sort_ratio(self.s9, self.s9a, full_process=True), 100)
        self.assertEqual(fuzz.partial_token_sort_ratio(self.s9, self.s9a, full_process=False), 100)
        self.assertEqual(fuzz.partial_token_sort_ratio(self.s10, self.s10a, full_process=False), 50)

    def testTokenSetRatio(self):
        self.assertEqual(fuzz.token_set_ratio(self.s4, self.s5), 100)
        self.assertEqual(fuzz.token_set_ratio(self.s8, self.s8a, full_process=False), 100)
        self.assertEqual(fuzz.token_set_ratio(self.s9, self.s9a, full_process=True), 100)
        self.assertEqual(fuzz.token_set_ratio(self.s9, self.s9a, full_process=False), 100)
        self.assertEqual(fuzz.token_set_ratio(self.s10, self.s10a, full_process=False), 50)

    def testPartialTokenSetRatio(self):
        self.assertEqual(fuzz.partial_token_set_ratio(self.s4, self.s7), 100)

    def testQuickRatioEqual(self):
        self.assertEqual(fuzz.QRatio(self.s1, self.s1a), 100)

    def testQuickRatioCaseInsensitive(self):
        self.assertEqual(fuzz.QRatio(self.s1, self.s2), 100)

    def testQuickRatioNotEqual(self):
        self.assertNotEqual(fuzz.QRatio(self.s1, self.s3), 100)

    def testWRatioEqual(self):
        self.assertEqual(fuzz.WRatio(self.s1, self.s1a), 100)

    def testWRatioCaseInsensitive(self):
        self.assertEqual(fuzz.WRatio(self.s1, self.s2), 100)

    def testWRatioPartialMatch(self):
        # a partial match is scaled by .9
        self.assertEqual(fuzz.WRatio(self.s1, self.s3), 90)

    def testWRatioMisorderedMatch(self):
        # misordered full matches are scaled by .95
        self.assertEqual(fuzz.WRatio(self.s4, self.s5), 95)

    def testWRatioUnicode(self):
        self.assertEqual(fuzz.WRatio(unicode(self.s1), unicode(self.s1a)), 100)

    def testQRatioUnicode(self):
        self.assertEqual(fuzz.WRatio(unicode(self.s1), unicode(self.s1a)), 100)

    def testEmptyStringsScore100(self):
        self.assertEqual(fuzz.ratio("", ""), 100)
        self.assertEqual(fuzz.partial_ratio("", ""), 100)

    def testIssueSeven(self):
        s1 = "HSINCHUANG"
        s2 = "SINJHUAN"
        s3 = "LSINJHUANG DISTRIC"
        s4 = "SINJHUANG DISTRICT"

        self.assertTrue(fuzz.partial_ratio(s1, s2) > 75)
        self.assertTrue(fuzz.partial_ratio(s1, s3) > 75)
        self.assertTrue(fuzz.partial_ratio(s1, s4) > 75)

    def testRatioUnicodeString(self):
        s1 = "\u00C1"
        s2 = "ABCD"
        score = fuzz.ratio(s1, s2)
        self.assertEqual(0, score)

    def testPartialRatioUnicodeString(self):
        s1 = "\u00C1"
        s2 = "ABCD"
        score = fuzz.partial_ratio(s1, s2)
        self.assertEqual(0, score)

    def testWRatioUnicodeString(self):
        s1 = "\u00C1"
        s2 = "ABCD"
        score = fuzz.WRatio(s1, s2)
        self.assertEqual(0, score)

        # Cyrillic.
        s1 = "\u043f\u0441\u0438\u0445\u043e\u043b\u043e\u0433"
        s2 = "\u043f\u0441\u0438\u0445\u043e\u0442\u0435\u0440\u0430\u043f\u0435\u0432\u0442"
        score = fuzz.WRatio(s1, s2, force_ascii=False)
        self.assertNotEqual(0, score)

        # Chinese.
        s1 = "\u6211\u4e86\u89e3\u6570\u5b66"
        s2 = "\u6211\u5b66\u6570\u5b66"
        score = fuzz.WRatio(s1, s2, force_ascii=False)
        self.assertNotEqual(0, score)

    def testQRatioUnicodeString(self):
        s1 = "\u00C1"
        s2 = "ABCD"
        score = fuzz.QRatio(s1, s2)
        self.assertEqual(0, score)

        # Cyrillic.
        s1 = "\u043f\u0441\u0438\u0445\u043e\u043b\u043e\u0433"
        s2 = "\u043f\u0441\u0438\u0445\u043e\u0442\u0435\u0440\u0430\u043f\u0435\u0432\u0442"
        score = fuzz.QRatio(s1, s2, force_ascii=False)
        self.assertNotEqual(0, score)

        # Chinese.
        s1 = "\u6211\u4e86\u89e3\u6570\u5b66"
        s2 = "\u6211\u5b66\u6570\u5b66"
        score = fuzz.QRatio(s1, s2, force_ascii=False)
        self.assertNotEqual(0, score)

    def testQratioForceAscii(self):
        s1 = "ABCD\u00C1"
        s2 = "ABCD"

        score = fuzz.QRatio(s1, s2, force_ascii=True)
        self.assertEqual(score, 100)

        score = fuzz.QRatio(s1, s2, force_ascii=False)
        self.assertLess(score, 100)

    def testQRatioForceAscii(self):
        s1 = "ABCD\u00C1"
        s2 = "ABCD"

        score = fuzz.WRatio(s1, s2, force_ascii=True)
        self.assertEqual(score, 100)

        score = fuzz.WRatio(s1, s2, force_ascii=False)
        self.assertLess(score, 100)

    def testTokenSetForceAscii(self):
        s1 = "ABCD\u00C1 HELP\u00C1"
        s2 = "ABCD HELP"

        score = fuzz._token_set(s1, s2, force_ascii=True)
        self.assertEqual(score, 100)

        score = fuzz._token_set(s1, s2, force_ascii=False)
        self.assertLess(score, 100)

    def testTokenSortForceAscii(self):
        s1 = "ABCD\u00C1 HELP\u00C1"
        s2 = "ABCD HELP"

        score = fuzz._token_sort(s1, s2, force_ascii=True)
        self.assertEqual(score, 100)

        score = fuzz._token_sort(s1, s2, force_ascii=False)
        self.assertLess(score, 100)
#########

@app.route('/verify',methods=['POST','GET'])
def verify():
    msg=""
    uname=""
    st=""
    cnt=0
    x=0
    f_arr=[]
    f_arr2=[]
    f_type=[]
    tt=100
        
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where id=%s",(pid,))
    d1 = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_patient_data")
    d2 = mycursor.fetchall()

    idd=0
    for dd in d2:

        if d1[7]==dd[6] and d1[8]==dd[7]:
            if d1[2]==dd[1] or d1[11]==dd[10]:
                idd=dd[0]
                break
        else:
            idd=0

    print("idd")
    print(idd)
    if idd>0:
        x+=1

        '''mycursor.execute("SELECT count(*) FROM mf_files where post_id=%s",(pid,))
        cnt = mycursor.fetchone()[0]
        
        mycursor.execute("SELECT * FROM mf_files where post_id=%s",(pid,))
        d3 = mycursor.fetchall()
        for d4 in d3:
            pfile=d4[3]
            file_ext=pfile.split(".")

            if file_ext[1]=="jpg" or file_ext[1]=="jpeg" or file_ext[1]=="png":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")

                    if fs[1]==file_ext[1]:
                        print(fname)
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        
                        per=tt-percentage_difference
                        print(per)
                        f_arr.append(pfile)
                        f_arr2.append(fname)
                        f_type.append("img")

                        if per<=98:
                            x+=1

            elif file_ext[1]=="docx":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")

                    if fs[1]==file_ext[1]:
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        per=tt-percentage_difference
                        f_arr.append(pfile)
                        f_arr2.append(fname)
                        f_type.append("docx")

                        if per<=98:                            
                            x+=1

            elif file_ext[1]=="pdf":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")

                    if fs[1]==file_ext[1]:
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        per=tt-percentage_difference
                        f_arr.append(pfile)
                        f_arr2.append(fname)
                        f_type.append("pdf")
                        if per<=98:                           
                            
                            x+=1'''
    if x>0:
        st="1"
        #mycursor.execute("update mf_post set req_status=1 where id=%s",(pid,))
        #mydb.commit()
    else:
        st="2"
        #mycursor.execute("update mf_post set req_status=2 where id=%s",(pid,))
        #mydb.commit()
         
       

    return render_template('web/verify.html',msg=msg,data=data,pid=pid,st=st,f_arr=f_arr,f_arr2=f_arr2,f_type=f_type)

def extract_text(file1):
    pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
    Actual_image = cv2.imread(file1)
    #Sample_img = cv2.resize(Actual_image,(400,350))
    Image_ht,Image_wd,Image_thickness = Actual_image.shape
    Sample_img = cv2.cvtColor(Actual_image,cv2.COLOR_BGR2RGB)
    texts = pytesseract.image_to_data(Sample_img) 
    mytext=""
    prevy=0

    
    
    for cnt,text in enumerate(texts.splitlines()):
        
        if cnt==0:
            continue
        text = text.split()
        if len(text)==12:
            x,y,w,h = int(text[6]),int(text[7]),int(text[8]),int(text[9])
            if(len(mytext)==0):
                prey=y
            if(prevy-y>=10 or y-prevy>=10):
                #print(mytext)
                s=1
                #mytext=""
            mytext = mytext + text[11]+"|"
            prevy=y

    mytext2=mytext.split("|")
    return mytext2


@app.route('/verify_file',methods=['POST','GET'])
def verify_file():
    msg=""
    uname=""
    d1=[]
    hd=[]
    st=""
    cnt=0
    x=0
    yx=0
    f_arr=[]
    f_arr2=[]
    f_type=[]
    f_cnt=[]
    data3=[]
    tt=100
    differences=""
    text1=""
    
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']

    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where id=%s",(pid,))
    d1 = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_patient_data")
    d2 = mycursor.fetchall()

    idd=0
    for dd in d2:

        if d1[7]==dd[6] and d1[8]==dd[7]:
            if d1[2]==dd[1] or d1[11]==dd[10]:
                idd=dd[0]
                break
        else:
            idd=0
        
    if idd>0:

        mycursor.execute("SELECT * FROM mf_patient_data where id=%s",(idd,))
        hd = mycursor.fetchone()

        mycursor.execute("SELECT count(*) FROM mf_files where post_id=%s",(pid,))
        cnt = mycursor.fetchone()[0]
        
        mycursor.execute("SELECT * FROM mf_files where post_id=%s",(pid,))
        d3 = mycursor.fetchall()
        
        for d4 in d3:
            pfile=d4[3]
            pcnt=d4[6]
            file_ext=pfile.split(".")

            if file_ext[1]=="jpg" or file_ext[1]=="jpeg" or file_ext[1]=="png":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")
                    
                    if fs[1]==file_ext[1]:
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        per=tt-percentage_difference
                        
                        f_arr.append(pfile)
                        f_cnt.append(pcnt)
                        f_arr2.append(fname)
                        f_type.append("img")
                        if per<=98:
                            x+=1

            elif file_ext[1]=="docx":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")

                    if fs[1]==file_ext[1]:
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        per=tt-percentage_difference
                        print(per)
                        f_arr.append(pfile)

                        f_cnt.append(pcnt)
                        f_arr2.append(fname)
                        f_type.append("docx")
                        if per<=98:
                            x+=1

            elif file_ext[1]=="pdf":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")

                    if fs[1]==file_ext[1]:
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        per=tt-percentage_difference
                        
                        f_arr.append(pfile)
                        f_cnt.append(pcnt)
                        f_arr2.append(fname)
                        f_type.append("pdf")
                        if per<=98:
                            x+=1
                    
        i=0
        for ft in f_type:
            dar=[]                   
            s=1
                            
            if ft=="img":
                ff1=""
                file1=f_arr[i]
                file2=f_arr2[i]
                
                dar.append(ft)
                dar.append(file1)
                dar.append(file2)
                i+=1
                print(file1)
                print(file2)
                
                ff1="static/upload/"+file1
                    
                before = cv2.imread(ff1)
                after = cv2.imread('static/data/'+file2)

                # Convert images to grayscale
                before_gray = cv2.cvtColor(before, cv2.COLOR_BGR2GRAY)
                after_gray = cv2.cvtColor(after, cv2.COLOR_BGR2GRAY)

                # Compute SSIM between the two images
                (score, diff) = structural_similarity(before_gray, after_gray, full=True)
                print("Image Similarity: {:.4f}%".format(score * 100))
                per=format(score * 100)

                # The diff image contains the actual image differences between the two images
                # and is represented as a floating point data type in the range [0,1] 
                # so we must convert the array to 8-bit unsigned integers in the range
                # [0,255] before we can use it with OpenCV
                diff = (diff * 255).astype("uint8")
                diff_box = cv2.merge([diff, diff, diff])

                # Threshold the difference image, followed by finding contours to
                # obtain the regions of the two input images that differ
                thresh = cv2.threshold(diff, 0, 255, cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)[1]
                contours = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                contours = contours[0] if len(contours) == 2 else contours[1]

                mask = np.zeros(before.shape, dtype='uint8')
                filled_after = after.copy()
                j=1
                for c in contours:
                    area = cv2.contourArea(c)
                    if area > 40:
                        x,y,w,h = cv2.boundingRect(c)
                        cv2.rectangle(before, (x, y), (x + w, y + h), (36,255,12), 2)
                        mm=cv2.rectangle(after, (x, y), (x + w, y + h), (36,255,12), 2)
                        cv2.imwrite("static/test/ggg.jpg", mm)

                        image = cv2.imread("static/test/ggg.jpg")
                        h1=h+10
                        w1=w+30
                        
                        
                        cropped = image[y:y+h1, x:x+w1]
                        gg="static/test/f"+str(j)+".jpg"
                        cv2.imwrite(""+gg, cropped)
                    
                        cv2.rectangle(diff_box, (x, y), (x + w, y + h), (36,255,12), 2)
                        cv2.drawContours(mask, [c], 0, (255,255,255), -1)
                        cv2.drawContours(filled_after, [c], 0, (0,255,0), -1)
                        j+=1
                print("###########")
                ###################
                mytext1=[]
                mytext2=[]
                pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
                Actual_image = cv2.imread("static/data/"+file2)
                #Sample_img = cv2.resize(Actual_image,(400,350))
                Image_ht,Image_wd,Image_thickness = Actual_image.shape
                Sample_img = cv2.cvtColor(Actual_image,cv2.COLOR_BGR2RGB)
                texts = pytesseract.image_to_data(Sample_img) 
                mytext=""
                prevy=0

                
                
                for cnt,text in enumerate(texts.splitlines()):
                    
                    if cnt==0:
                        continue
                    text = text.split()
                    if len(text)==12:
                        x,y,w,h = int(text[6]),int(text[7]),int(text[8]),int(text[9])
                        if(len(mytext)==0):
                            prey=y
                        if(prevy-y>=10 or y-prevy>=10):
                            #print(mytext)
                            s=1
                            #mytext=""
                        mytext = mytext + text[11]+"|"
                        prevy=y

                mytext1=mytext.split("|")
                dar.append(mytext1)
                ################
                pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
                Actual_image = cv2.imread("static/upload/"+file1)
                #Sample_img = cv2.resize(Actual_image,(400,350))
                Image_ht,Image_wd,Image_thickness = Actual_image.shape
                Sample_img = cv2.cvtColor(Actual_image,cv2.COLOR_BGR2RGB)
                texts = pytesseract.image_to_data(Sample_img) 
                mytext=""
                prevy=0

                
                
                for cnt,text in enumerate(texts.splitlines()):
                    
                    if cnt==0:
                        continue
                    text = text.split()
                    if len(text)==12:
                        x,y,w,h = int(text[6]),int(text[7]),int(text[8]),int(text[9])
                        if(len(mytext)==0):
                            prey=y
                        if(prevy-y>=10 or y-prevy>=10):
                            #print(mytext)
                            s=1
                            #mytext=""
                        mytext = mytext + text[11]+"|"
                        prevy=y

                mytext2=mytext.split("|")
                m=0
                dyy=[]
               
                for mt in mytext2:
                    dy=[]
                    if mt==mytext1[m]:
                        dy.append(mt)
                        dy.append("1")
                    else:
                        yx+=1
                        dy.append(mt)
                        dy.append("2")
                        
                    m+=1
                    dyy.append(dy)
                    
                dar.append(dyy)

                
                
            ###############################################################
            elif ft=="docx":
                ff1=""
                file1=f_arr[i]
                file2=f_arr2[i]
                fcnt=f_cnt[i]

                ni=1
                f2=file2.split(".")
                while ni<=fcnt:
                    fn="m"+str(pid)+"_"+str(ni)+".png"
                    file1=fn

                    fn2=f2[0]+"_"+str(ni)+".png"
                    file2=fn2
                    ni+=1

                dar.append(ft)
                dar.append(file1)
                dar.append(file2)
                mytext1=extract_text("static/data/"+file2)
                dar.append(mytext1)

                mytext2=extract_text("static/upload/"+file1)
                m=0
                dyy=[]
                
                for mt in mytext2:
                    dy=[]
                    if mt==mytext1[m]:
                        dy.append(mt)
                        dy.append("1")
                    else:
                        yx+=1
                        dy.append(mt)
                        dy.append("2")
                        
                    m+=1
                    dyy.append(dy)
                    
                dar.append(dyy)
                i+=1

               
                '''dar.append(ft)
                dar.append(file1)
                dar.append(file2)
                i+=1
                text1 = extract_text_from_docx("static/data/"+file2)
                text2 = extract_text_from_docx("static/upload/"+file1)

                differences = highlight_differences(text1, text2)
                print("docx diff")
                dar.append(text1)
                dar.append(differences)'''


                

            elif ft=="pdf":
               
                ff1=""
                file1=f_arr[i]
                file2=f_arr2[i]
                fcnt=f_cnt[i]
                
                f1=file1.split(".")
                file11=f1[0]+".docx"

                f2=file2.split(".")
                file22=f2[0]+".docx"

                ni=1
                
                while ni<=fcnt:
                    fn="m"+str(pid)+"_"+str(ni)+".png"
                    file1=fn

                    fn2=f2[0]+"_"+str(ni)+".png"
                    file2=fn2
                    ni+=1

                dar.append(ft)
                dar.append(file1)
                dar.append(file2)
                mytext1=extract_text("static/data/"+file2)
                dar.append(mytext1)

                mytext2=extract_text("static/upload/"+file1)
                m=0
                dyy=[]
                
                for mt in mytext2:
                    dy=[]
                    if mt==mytext1[m]:
                        dy.append(mt)
                        dy.append("1")
                    else:
                        yx+=1
                        dy.append(mt)
                        dy.append("2")
                        
                    m+=1
                    dyy.append(dy)
                    
                dar.append(dyy)
                i+=1

                
                '''text1 = extract_text_from_docx("static/data/"+file22)
                text2 = extract_text_from_docx("static/upload/"+file11)

                differences = highlight_differences(text1, text2)'''

            else:
                s=1
                
            ################################################################    
            data3.append(dar)                        
            
        if yx==0:
            st="1"
            mycursor.execute("update mf_post set req_status=1 where id=%s",(pid,))
            mydb.commit()
        else:
            st="2"
            mycursor.execute("update mf_post set req_status=2 where id=%s",(pid,))
            mydb.commit()
        
    return render_template('web/verify_file.html',msg=msg,data=data,hd=hd,d1=d1,pid=pid,st=st,f_arr=f_arr,f_arr2=f_arr2,f_type=f_type,data3=data3)

@app.route('/admin',methods=['POST','GET'])
def admin():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user")
    data = mycursor.fetchall()

  
    return render_template('web/admin.html',msg=msg,data=data)

@app.route('/admin2',methods=['POST','GET'])
def admin2():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_donator")
    data = mycursor.fetchall()

  
    return render_template('web/admin2.html',msg=msg,data=data)



@app.route('/user_post',methods=['POST','GET'])
def user_post():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where uname=%s order by id desc",(uname,))
    data2 = mycursor.fetchall()
    
    
    return render_template('web/user_post.html',msg=msg,data=data,data2=data2)

@app.route('/user_donate',methods=['POST','GET'])
def user_donate():
    msg=""
    uname=""
    st=""
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT count(*) FROM mf_donation where pid=%s order by id desc",(pid,))
    d1 = mycursor.fetchone()[0]
    if d1>0:
        st="1"
    
    mycursor.execute("SELECT * FROM mf_donation where pid=%s order by id desc",(pid,))
    data2 = mycursor.fetchall()
    
    
    return render_template('web/user_donate.html',msg=msg,data=data,data2=data2,st=st)

@app.route('/user_fund',methods=['POST','GET'])
def user_fund():
    msg=""
    uname=""
    st=""
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT count(*) FROM mf_donation where uname=%s order by id desc",(uname,))
    d1 = mycursor.fetchone()[0]
    if d1>0:
        st="1"
    
    mycursor.execute("SELECT * FROM mf_donation where uname=%s order by id desc",(uname,))
    data2 = mycursor.fetchall()
    
    
    return render_template('web/user_fund.html',msg=msg,data=data,data2=data2,st=st)

@app.route('/dr_home',methods=['POST','GET'])
def dr_home():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_donator where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post p,mf_user u where p.uname=u.uname && p.req_status=1 && p.fund_amount=0 order by p.id desc")
    data2 = mycursor.fetchall()
    
    
    return render_template('web/dr_home.html',msg=msg,data=data,data2=data2)

@app.route('/dr_donate',methods=['POST','GET'])
def dr_donate():
    msg=""
    uname=""
    st=""
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_donator where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where id=%s",(pid,))
    data2 = mycursor.fetchone()
    user=data2[1]
    famt=data2[17]

    now = datetime.datetime.now()
    rdate=now.strftime("%d-%m-%Y")

    if request.method=='POST':
        amount=request.form['amount']
        trans_id = request.form['trans_id']

        mycursor.execute("SELECT max(id)+1 FROM mf_donation")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1
            
        amt=famt+int(amount)
        mycursor.execute("update mf_post set fund_amount=%s where id=%s",(amt,pid))
        mydb.commit()
        
         
        sql = "INSERT INTO mf_donation(id,donator,uname,pid,amount,trans_id,rdate) VALUES (%s, %s, %s, %s, %s,%s,%s)"
        val = (maxid,uname,user,pid,amount,trans_id,rdate)
        mycursor.execute(sql, val)
        mydb.commit()
        msg="success"

    mycursor.execute("SELECT count(*) FROM mf_donation where pid=%s",(pid,))
    d1 = mycursor.fetchone()[0]
    if d1>0:
        st="1"
    
    mycursor.execute("SELECT * FROM mf_donation where pid=%s",(pid,))
    data2 = mycursor.fetchall()
    
    return render_template('web/dr_donate.html',msg=msg,data=data,data2=data2,st=st)


@app.route('/dr_post',methods=['POST','GET'])
def dr_post():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_donator where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where req_status=1 order by id desc")
    data2 = mycursor.fetchall()
    
    
    return render_template('web/dr_post.html',msg=msg,data=data,data2=data2)

@app.route('/dr_bill', methods=['GET', 'POST'])
def dr_bill():
    msg=""
    st=""
    data2=[]
    view=""
    fname=""
    textpdf=""
    textdoc=""
    ccode=""
    pid=request.args.get("pid")
    nn=0
   
    ss='0'
    if pid=='0':
        ss='0'
    else:
        ss='1'
        
    mycursor = mydb.cursor()
 

    mycursor.execute("SELECT * FROM mf_files where post_id=%s",(pid,))
    data1 = mycursor.fetchall()

    mycursor.execute("SELECT count(*) FROM mf_files where post_id=%s",(pid,))
    cnt = mycursor.fetchone()[0]
    if cnt>0:
        st="1"
        
        mycursor.execute("SELECT * FROM mf_files where post_id=%s",(pid,))
        d1 = mycursor.fetchall()
        for d2 in d1:
            dt=[]
            dt.append(d2[0])
            dt.append(d2[1])
            dt.append(d2[2])
            dt.append(d2[3])
            dt.append(d2[4])
            dt.append(d2[5])
            dt.append(d2[6])

            #7#
            ex=d2[3].split(".")
            if ex[1]=="png":
                dt.append("png")
            elif ex[1]=="jpg":
                dt.append("jpg")
            elif ex[1]=="jpeg":
                dt.append("jpeg")
            elif ex[1]=="pdf":
                dt.append("pdf")
                fname=d2[3]
                file2="static/upload/"+fname
                textpdf = extract_text_from_pdf(file2)
               
            
            elif ex[1]=="docx":
                dt.append("docx")
                fname=d2[3]
                file2="static/upload/"+fname
    
                textdoc = extract_text_from_docx(file2)
                
            
            elif ex[1]=="txt":
                dt.append("txt")
                fname=d2[3]
                file1 = open("static/upload/"+fname, 'r')
                Lines = file1.readlines()
                 
                count = 0
                result=""
                # Strips the newline character
                for line in Lines:
                    result = "".join(line for line in Lines if not line.isspace())
                    count += 1
                    #print("Line{}: {}".format(count, line.strip()))
                ccode=result
            
            else:
                dt.append("")

            #8#
            if d2[6]>0:
                dt.append("yes")
            else:
                dt.append("no")
            #9#
            if d2[6]>0:
                i=1
                dtt=[]
                
                while i<=d2[6]:
                    mg="m"+str(d2[0])+"_"+str(i)+".png"
                    dtt.append(mg)
                    i+=1
                dt.append(dtt)
            ##
            data2.append(dt)
           

    return render_template('web/dr_bill.html',msg=msg,data2=data2,st=st)

@app.route('/dr_fund',methods=['POST','GET'])
def dr_fund():
    msg=""
    uname=""
    st=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_donator where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT count(*) FROM mf_donation where donator=%s order by id desc",(uname,))
    d1 = mycursor.fetchone()[0]
    if d1>0:
        st="1"
    
    mycursor.execute("SELECT * FROM mf_donation where donator=%s order by id desc",(uname,))
    data2 = mycursor.fetchall()
    
    
    return render_template('web/dr_fund.html',msg=msg,data=data,data2=data2,st=st)


@app.route('/logout')
def logout():
    # remove the username from the session if it is there
    #session.pop('username', None)
    return redirect(url_for('index'))

if __name__ == "__main__":
    app.secret_key = os.urandom(12)
    app.run(debug=True,host='0.0.0.0', port=5000)
